import logging
from pathlib import Path
from typing import Optional, List
from moviepy import VideoFileClip
from pydub import AudioSegment
from openai import OpenAI
from azure.storage.blob import BlobServiceClient
from azure.core.exceptions import ResourceExistsError
from docx import Document
from concurrent.futures import ThreadPoolExecutor
from functools import partial
import time
import math
import os
import asyncio
from utils.cleanup import split_into_sentences
import io

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)

# Suppress Azure SDK logging
logging.getLogger('azure').setLevel(logging.WARNING)
logging.getLogger('azure.core.pipeline.policies.http_logging_policy').setLevel(logging.WARNING)

logger = logging.getLogger(__name__)

class VideoProcessor:
    def __init__(self):
        self.blob_service_client = self._get_blob_service_client()
        self.containers = self._ensure_containers()
        # Create a thread pool for CPU-bound tasks
        self.executor = ThreadPoolExecutor(max_workers=3)

    def _get_blob_service_client(self):
        """Initialize Azure Blob Service Client."""
        connection_string = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
        if not connection_string:
            raise ValueError("Azure Storage connection string not found")
        return BlobServiceClient.from_connection_string(connection_string)

    def _ensure_containers(self):
        """Ensure required containers exist."""
        containers = {
            'videos': 'video-uploads',
            'transcripts': 'transcriptions'
        }
        
        for container_name in containers.values():
            try:
                self.blob_service_client.create_container(container_name)
                logger.info(f"Created container: {container_name}")
            except ResourceExistsError:
                logger.debug(f"Container already exists: {container_name}")
        
        return containers

    @staticmethod
    def extract_audio(video_path, output_path):
        """Extract audio from video file and save as MP3."""
        try:
            video = VideoFileClip(video_path)
            try:
                audio = video.audio
                audio.write_audiofile(output_path)
                return True
            finally:
                if hasattr(video, 'close'):
                    video.close()
                if hasattr(video.audio, 'close'):
                    video.audio.close()
        except Exception as e:
            logger.error(f"Error extracting audio: {e}")
            return False

    @staticmethod
    def split_audio(audio_path):
        """Split audio file into 10-minute chunks."""
        try:
            audio = AudioSegment.from_mp3(audio_path)
            chunk_length = 10 * 60 * 1000  # 10 minutes in milliseconds
            chunks = []
            
            for i in range(0, len(audio), chunk_length):
                chunk = audio[i:i + chunk_length]
                chunk_path = f"{audio_path.stem}_chunk_{i//chunk_length}{audio_path.suffix}"
                chunk_path = audio_path.parent / chunk_path
                chunk.export(str(chunk_path), format="mp3")
                chunks.append(chunk_path)
                
            return chunks
        except Exception as e:
            logger.error(f"Error splitting audio: {e}")
            return []

    async def process_video(self, video_name: str, client: OpenAI) -> bool:
        """Process a single video file from Azure Storage."""
        temp_files = []
        try:
            start_time = time.time()
            
            # Setup paths
            temp_dir = Path("temp")
            temp_dir.mkdir(exist_ok=True)
            temp_video_path = temp_dir / video_name
            temp_audio_path = temp_dir / f"{video_name}.mp3"
            temp_files.extend([temp_video_path, temp_audio_path])

            # Download video
            download_start = time.time()
            logger.info(f"Starting download of video: {video_name}")
            blob_client = self.blob_service_client.get_container_client(
                self.containers['videos']).get_blob_client(video_name)
            
            with open(temp_video_path, "wb") as video_file:
                download_stream = await asyncio.get_event_loop().run_in_executor(
                    self.executor,
                    blob_client.download_blob
                )
                await asyncio.get_event_loop().run_in_executor(
                    self.executor,
                    lambda: download_stream.readinto(video_file)
                )
            logger.info(f"Download took {time.time() - download_start:.2f} seconds")

            # Extract audio
            extract_start = time.time()
            logger.info("Starting audio extraction")
            video = VideoFileClip(str(temp_video_path))
            video.audio.write_audiofile(str(temp_audio_path))
            video.close()
            logger.info(f"Audio extraction took {time.time() - extract_start:.2f} seconds")

            # Split and transcribe
            split_start = time.time()
            audio = AudioSegment.from_mp3(temp_audio_path)
            chunk_length = 5 * 60 * 1000  # 5 minutes
            chunks = []
            
            # Create chunks
            for i in range(0, len(audio), chunk_length):
                chunk_path = temp_dir / f"chunk_{i//chunk_length}.mp3"
                audio[i:i + chunk_length].export(chunk_path, format="mp3")
                chunks.append(chunk_path)
                temp_files.append(chunk_path)
            
            logger.info(f"Split into {len(chunks)} chunks")

            # Transcribe chunks in parallel
            sem = asyncio.Semaphore(10)  # Increased concurrent API calls
            
            async def transcribe_chunk(chunk_path):
                async with sem:
                    return await self.transcribe_with_retry(chunk_path, client)

            # Create and run tasks in parallel
            transcribe_tasks = [transcribe_chunk(chunk) for chunk in chunks]
            logger.info(f"Starting parallel transcription of {len(chunks)} chunks")
            results = await asyncio.gather(*transcribe_tasks)
            logger.info("Completed all transcriptions")
            full_transcript = [t for t in results if t is not None]

            # Create and upload document
            if full_transcript:
                doc = Document()
                text = '\n'.join(full_transcript)
                sentences = split_into_sentences(text)
                for sentence in sentences:
                    doc.add_paragraph(sentence)

                output_filename = f"{video_name}_cleaned.docx"
                doc.save(str(temp_dir / output_filename))
                
                with open(temp_dir / output_filename, 'rb') as doc_file:
                    await asyncio.get_event_loop().run_in_executor(
                        self.executor,
                        lambda: self.blob_service_client.get_container_client(
                            self.containers['transcripts']
                        ).upload_blob(output_filename, doc_file, overwrite=True)
                    )

            # Cleanup
            for temp_file in temp_files:
                if temp_file.exists():
                    temp_file.unlink()

            total_time = time.time() - start_time
            logger.info(f"Total processing time: {total_time:.2f} seconds")
            return True

        except Exception as e:
            logger.error(f"Error processing video: {e}")
            # Cleanup on error
            for temp_file in temp_files:
                if isinstance(temp_file, Path) and temp_file.exists():
                    temp_file.unlink()
            return False

    async def _cleanup_files(self, temp_files, video_name=None):
        """Clean up temporary files and blobs."""
        for temp_file in temp_files:
            try:
                if isinstance(temp_file, Path) and temp_file.exists():
                    await asyncio.get_event_loop().run_in_executor(
                        self.executor,
                        temp_file.unlink
                    )
            except Exception as e:
                logger.error(f"Error cleaning up {temp_file}: {e}")

        if video_name:
            try:
                await asyncio.get_event_loop().run_in_executor(
                    self.executor,
                    lambda: self.blob_service_client.get_container_client(self.containers['video-uploads']).delete_blob(video_name)
                )
            except Exception as e:
                logger.error(f"Error cleaning up blob: {e}")

    def _create_document(self, transcript_list):
        """Create a Word document from transcripts"""
        doc = Document()
        text = '\n'.join(transcript_list)
        sentences = split_into_sentences(text)
        for sentence in sentences:
            doc.add_paragraph(sentence)
        return doc

    async def transcribe_with_retry(self, audio_path: Path, client: OpenAI, max_retries: int = 3) -> Optional[str]:
        """Async version of transcribe with retry"""
        async def _transcribe():
            logger.info(f"Starting transcription of {audio_path}")
            # Run the OpenAI API call in the thread pool
            result = await asyncio.get_event_loop().run_in_executor(
                self.executor,
                lambda: client.audio.transcriptions.create(
                    model="whisper-1",
                    file=open(audio_path, "rb"),
                    response_format="text"
                )
            )
            logger.info(f"Completed transcription of {audio_path}")
            return result

        for attempt in range(max_retries):
            try:
                transcript = await _transcribe()
                return transcript
            except Exception as e:
                logger.warning(f"Attempt {attempt + 1} failed for {audio_path}: {e}")
                if attempt == max_retries - 1:
                    logger.error(f"Final retry failed for {audio_path}: {e}")
                    return None
                await asyncio.sleep(math.pow(2, attempt))

    def list_videos(self) -> List[str]:
        """List all MP4 files in the videos container."""
        container_client = self.blob_service_client.get_container_client(self.containers['videos'])
        return [blob.name for blob in container_client.list_blobs() if blob.name.lower().endswith('.mp4')]

# Add this test section at the end of the file
if __name__ == "__main__":
    import os
    from dotenv import load_dotenv
    import asyncio
    
    # Load environment variables
    load_dotenv()
    
    # Initialize OpenAI client
    client = OpenAI()
    
    # Initialize video processor
    processor = VideoProcessor()
    
    async def test_video_processing():
        '''Test the video processing pipeline'''
        try:
            # List available videos
            print("Available videos:")
            videos = processor.list_videos()
            for i, video in enumerate(videos, 1):
                print(f"{i}. {video}")
            
            if not videos:
                print("No videos found in storage.")
                return
            
            # Let user select a video
            while True:
                try:
                    choice = int(input("\nEnter the number of the video to process (0 to exit): ")) - 1
                    if choice == -1:
                        return
                    if 0 <= choice < len(videos):
                        break
                    print("Invalid choice. Please try again.")
                except ValueError:
                    print("Please enter a valid number.")
            
            video_name = videos[choice]
            print(f"\nProcessing video: {video_name}")
            
            # Process the video - now with await
            success = await processor.process_video(video_name, client)
            
            if success:
                print(f"\nSuccessfully processed {video_name}")
                print("Check the 'transcripts' container in Azure Storage for the output.")
            else:
                print(f"\nFailed to process {video_name}")
                
        except Exception as e:
            print(f"Error during testing: {e}")
    
    # Run the test with asyncio
    print("Video Processing Test Tool")
    print("=========================")
    asyncio.run(test_video_processing()) 

